# rag_engine.py
import os
import logging
import requests
import tempfile
import json
from bs4 import BeautifulSoup
from pathlib import Path
from urllib.parse import urljoin
import PyPDF2
import docx

from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.llms import Ollama
from langchain.chains import RetrievalQA
from langchain.schema import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger('rag_engine')

class RAGEngine:
    """Retrieval Augmented Generation engine for ADGM documents"""
    
    def __init__(self, model_name="llama3:8b", data_dir="adgm_extracted_data"):
        logger.info("Initializing RAG Engine")
        self.data_dir = data_dir
        self.adgm_sources = self._get_adgm_sources()
        
        # Initialize embedding model
        try:
            from langchain.embeddings import HuggingFaceEmbeddings
            self.embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
            logger.info("Using HuggingFace embeddings")
        except ImportError:
            logger.warning("HuggingFace embeddings not available, using fallback")
            from langchain.embeddings import FakeEmbeddings
            self.embeddings = FakeEmbeddings(size=768)
            logger.info("Using fallback embeddings")
        
        # Initialize with Ollama model
        try:
            self.llm = Ollama(model=model_name)
            logger.info(f"Successfully initialized Ollama with model: {model_name}")
        except Exception as e:
            logger.error(f"Error initializing Ollama: {str(e)}")
            self.llm = None
        
        self.vectorstore = None
        self.retriever = None
        self.qa_chain = None
        
        # Try to load existing vector store or create a new one
        self.load_or_create_vectorstore()
    
    def _get_adgm_sources(self):
        """Define ADGM data sources"""
        return [
            ("Company Formation", "General Incorporation", "https://www.adgm.com/registration-authority/registration-and-incorporation"),
            ("Company Formation", "Resolution", "https://assets.adgm.com/download/assets/adgm-ra-resolution-multiple-incorporate-shareholders-LTD-incorporation-v2.docx/186a12846c3911efa4e6c6223862cd87"),
            ("Company Formation", "Templates", "https://www.adgm.com/setting-up"),
            ("Policy & Guidance", "Templates", "https://www.adgm.com/legal-framework/guidance-and-policy-statements"),
            ("Company Setup", "Checklist", "https://www.adgm.com/documents/registration-authority/registration-and-incorporation/checklist/branch-non-financial-services-20231228.pdf"),
            ("Company Setup", "Private Company", "https://www.adgm.com/documents/registration-authority/registration-and-incorporation/checklist/private-company-limited-by-guarantee-non-financial-services-20231228.pdf"),
            ("Employment", "Contract 2024", "https://assets.adgm.com/download/assets/ADGM+Standard+Employment+Contract+Template+-+ER+2024+(Feb+2025).docx/ee14b252edbe11efa63b12b3a30e5e3a"),
            ("Employment", "Contract 2019", "https://assets.adgm.com/download/assets/ADGM+Standard+Employment+Contract+-+ER+2019+-+Short+Version+(May+2024).docx/33b57a92ecfe11ef97a536cc36767ef8"),
            ("Data Protection", "Policy", "https://www.adgm.com/documents/office-of-data-protection/templates/adgm-dpr-2021-appropriate-policy-document.pdf"),
            ("Compliance", "Annual Accounts", "https://www.adgm.com/operating-in-adgm/obligations-of-adgm-registered-entities/annual-filings/annual-accounts"),
            ("Permits", "Application", "https://www.adgm.com/operating-in-adgm/post-registration-services/letters-and-permits"),
            ("Regulatory", "Incorporation", "https://en.adgm.thomsonreuters.com/rulebook/7-company-incorporation-package"),
            ("Regulatory", "Shareholder Resolution", "https://assets.adgm.com/download/assets/Templates_SHReso_AmendmentArticles-v1-20220107.docx/97120d7c5af911efae4b1e183375c0b2?forcedownload=1"),
        ]
    
    def load_or_create_vectorstore(self):
        """Load existing vector store or create a new one from documents"""
        vectorstore_path = os.path.join(self.data_dir, "vectorstore")
        
        # Try to load existing vector store
        if os.path.exists(vectorstore_path) and os.path.isfile(os.path.join(vectorstore_path, "index.faiss")):
            try:
                logger.info(f"Loading existing vector store from {vectorstore_path}")
                self.vectorstore = FAISS.load_local(vectorstore_path, self.embeddings)
                logger.info(f"Vector store loaded successfully")
                self._setup_retriever()
                return
            except Exception as e:
                logger.error(f"Error loading vector store: {str(e)}")
        
        # Create new vector store from documents
        logger.info("Existing vector store not found or could not be loaded")
        logger.info("Creating new vector store from ADGM sources")
        self.collect_and_create_vectorstore()
    
    def collect_and_create_vectorstore(self):
        """Collect data from ADGM sources and create vectorstore"""
        logger.info("Starting collection of ADGM data")
        
        # Create directories if they don't exist
        os.makedirs(self.data_dir, exist_ok=True)
        documents_dir = os.path.join(self.data_dir, "documents")
        extracted_text_dir = os.path.join(self.data_dir, "extracted_text")
        os.makedirs(documents_dir, exist_ok=True)
        os.makedirs(extracted_text_dir, exist_ok=True)
        
        # Headers for HTTP requests
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        # Store documents and metadata
        documents = []
        metadata_list = []
        
        # Process each source
        for category, doc_type, url in self.adgm_sources:
            try:
                logger.info(f"Processing {category} - {doc_type} from {url}")
                
                # Handle different URL types
                if url.lower().endswith('.pdf'):
                    # Download and process PDF
                    response = requests.get(url, headers=headers, timeout=30)
                    response.raise_for_status()
                    
                    # Save PDF
                    pdf_filename = f"{category}_{doc_type}.pdf"
                    pdf_path = os.path.join(documents_dir, pdf_filename)
                    with open(pdf_path, 'wb') as f:
                        f.write(response.content)
                    
                    # Extract text
                    pdf_text = self._extract_text_from_pdf(pdf_path)
                    
                    if pdf_text:
                        # Save extracted text
                        text_path = os.path.join(extracted_text_dir, f"{category}_{doc_type}.txt")
                        with open(text_path, 'w', encoding='utf-8') as f:
                            f.write(pdf_text)
                        
                        # Add to documents
                        documents.append(Document(
                            page_content=pdf_text,
                            metadata={
                                "source": url,
                                "category": category,
                                "document_type": doc_type,
                                "file_type": "PDF"
                            }
                        ))
                        
                        # Add to metadata
                        metadata_list.append({
                            "source": url,
                            "category": category,
                            "document_type": doc_type,
                            "filename": pdf_filename,
                            "file_type": "PDF"
                        })
                
                elif url.lower().endswith('.docx'):
                    # Download and process DOCX
                    response = requests.get(url, headers=headers, timeout=30)
                    response.raise_for_status()
                    
                    # Save DOCX
                    docx_filename = f"{category}_{doc_type}.docx"
                    docx_path = os.path.join(documents_dir, docx_filename)
                    with open(docx_path, 'wb') as f:
                        f.write(response.content)
                    
                    # Extract text
                    docx_text = self._extract_text_from_docx(docx_path)
                    
                    if docx_text:
                        # Save extracted text
                        text_path = os.path.join(extracted_text_dir, f"{category}_{doc_type}.txt")
                        with open(text_path, 'w', encoding='utf-8') as f:
                            f.write(docx_text)
                        
                        # Add to documents
                        documents.append(Document(
                            page_content=docx_text,
                            metadata={
                                "source": url,
                                "category": category,
                                "document_type": doc_type,
                                "file_type": "DOCX"
                            }
                        ))
                        
                        # Add to metadata
                        metadata_list.append({
                            "source": url,
                            "category": category,
                            "document_type": doc_type,
                            "filename": docx_filename,
                            "file_type": "DOCX"
                        })
                
                else:
                    # Process webpage
                    response = requests.get(url, headers=headers, timeout=30)
                    response.raise_for_status()
                    
                    # Save HTML
                    html_filename = f"{category}_{doc_type}.html"
                    html_path = os.path.join(documents_dir, html_filename)
                    with open(html_path, 'w', encoding='utf-8') as f:
                        f.write(response.text)
                    
                    # Parse HTML
                    soup = BeautifulSoup(response.text, 'html.parser')
                    
                    # Extract main text content
                    for script in soup(["script", "style", "header", "footer", "nav"]):
                        script.extract()
                    
                    webpage_text = soup.get_text(separator=' ', strip=True)
                    
                    # Save extracted text
                    text_path = os.path.join(extracted_text_dir, f"{category}_{doc_type}.txt")
                    with open(text_path, 'w', encoding='utf-8') as f:
                        f.write(webpage_text)
                    
                    # Add to documents
                    documents.append(Document(
                        page_content=webpage_text,
                        metadata={
                            "source": url,
                            "category": category,
                            "document_type": doc_type,
                            "file_type": "Webpage"
                        }
                    ))
                    
                    # Add to metadata
                    metadata_list.append({
                        "source": url,
                        "category": category,
                        "document_type": doc_type,
                        "filename": html_filename,
                        "file_type": "Webpage"
                    })
                    
                    # Find and process document links in the webpage
                    self._extract_embedded_documents(soup, url, category, doc_type, documents_dir, extracted_text_dir, documents, metadata_list)
                
            except Exception as e:
                logger.error(f"Error processing {url}: {str(e)}")
        
        # Save metadata
        metadata_path = os.path.join(self.data_dir, "document_metadata.json")
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata_list, f, indent=2)
        
        # Create vector store if documents were collected
        if documents:
            logger.info(f"Creating vector store from {len(documents)} documents")
            
            # Split documents into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200
            )
            
            chunks = text_splitter.split_documents(documents)
            logger.info(f"Split into {len(chunks)} chunks")
            
            # Create vector store
            self.vectorstore = FAISS.from_documents(chunks, self.embeddings)
            
            # Save vector store
            vectorstore_path = os.path.join(self.data_dir, "vectorstore")
            os.makedirs(vectorstore_path, exist_ok=True)
            self.vectorstore.save_local(vectorstore_path)
            
            logger.info(f"Vector store created and saved to {vectorstore_path}")
            
            # Set up retriever
            self._setup_retriever()
        else:
            logger.error("No documents collected, vector store creation failed")
    
    def _extract_embedded_documents(self, soup, base_url, category, doc_type, documents_dir, extracted_text_dir, documents, metadata_list):
        """Extract and process documents embedded in a webpage"""
        # Find all links
        links = soup.find_all('a', href=True)
        
        for link in links:
            href = link['href']
            
            # Skip empty links or javascript
            if not href or href.startswith('javascript:') or href == '#':
                continue
            
            # Make absolute URL if relative
            if not href.startswith('http'):
                href = urljoin(base_url, href)
            
            # Check if it's a document link (.pdf or .docx)
            if href.lower().endswith('.pdf') or href.lower().endswith('.docx'):
                try:
                    logger.info(f"Found embedded document: {href}")
                    
                    # Get link text for better metadata
                    link_text = link.get_text().strip()
                    if not link_text:
                        link_text = "Embedded Document"
                    
                    # Determine file type
                    file_type = "PDF" if href.lower().endswith('.pdf') else "DOCX"
                    
                    # Download the document
                    response = requests.get(href, headers={'User-Agent': 'Mozilla/5.0'}, timeout=30)
                    response.raise_for_status()
                    
                    # Save to temporary file
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_type.lower()}") as temp_file:
                        temp_file.write(response.content)
                        temp_path = temp_file.name
                    
                    # Extract text
                    if file_type == "PDF":
                        extracted_text = self._extract_text_from_pdf(temp_path)
                    else:
                        extracted_text = self._extract_text_from_docx(temp_path)
                    
                    # Save the document and extracted text
                    if extracted_text:
                        # Create safe filename
                        safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in link_text)
                        safe_name = safe_name[:50]  # Limit length
                        
                        # Save document
                        doc_filename = f"{category}_{doc_type}_{safe_name}.{file_type.lower()}"
                        doc_path = os.path.join(documents_dir, doc_filename)
                        with open(doc_path, 'wb') as f:
                            f.write(response.content)
                        
                        # Save extracted text
                        text_filename = f"{category}_{doc_type}_{safe_name}.txt"
                        text_path = os.path.join(extracted_text_dir, text_filename)
                        with open(text_path, 'w', encoding='utf-8') as f:
                            f.write(extracted_text)
                        
                        # Add to documents
                        documents.append(Document(
                            page_content=extracted_text,
                            metadata={
                                "source": href,
                                "category": category,
                                "document_type": f"{doc_type} - {link_text}",
                                "file_type": file_type,
                                "parent_url": base_url
                            }
                        ))
                        
                        # Add to metadata
                        metadata_list.append({
                            "source": href,
                            "category": category,
                            "document_type": f"{doc_type} - {link_text}",
                            "filename": doc_filename,
                            "text_path": text_path,
                            "file_type": file_type,
                            "parent_url": base_url
                        })
                    
                    # Clean up temp file
                    os.unlink(temp_path)
                    
                except Exception as e:
                    logger.error(f"Error processing embedded document {href}: {str(e)}")
    
    def _extract_text_from_pdf(self, pdf_path):
        """Extract text from a PDF file"""
        try:
            text = ""
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    text += pdf_reader.pages[page_num].extract_text() + "\n\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
            return None
    
    def _extract_text_from_docx(self, docx_path):
        """Extract text from a DOCX file"""
        try:
            doc = docx.Document(docx_path)
            text = ""
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
            return None
    
    def _setup_retriever(self):
        """Set up retriever and QA chain"""
        if not self.vectorstore or not self.llm:
            logger.error("Cannot set up retriever: vectorstore or LLM not available")
            return
        
        self.retriever = self.vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 5}  # Retrieve top 5 most relevant chunks
        )
        
        self.qa_chain = RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=self.retriever,
            return_source_documents=True
        )
        
        logger.info("Retriever and QA chain set up successfully")
    
    def query(self, question):
        """Query the RAG system"""
        if not self.vectorstore or not self.qa_chain or not self.llm:
            logger.warning("RAG system not fully initialized, using fallback")
            return "RAG system not fully initialized. Please ensure the vector database is properly built."
        
        try:
            logger.info(f"Processing query: {question}")
            
            # Add ADGM context to the query
            enhanced_query = f"""
            Based on ADGM (Abu Dhabi Global Market) regulations and templates, please answer the following:
            
            {question}
            
            Focus specifically on ADGM requirements, regulations, and templates.
            """
            
            result = self.qa_chain({"query": enhanced_query})
            
            # Extract answer and sources
            answer = result.get("result", "")
            source_docs = result.get("source_documents", [])
            
            # Add source information
            if source_docs:
                sources_text = "\n\nSources:\n"
                unique_sources = set()
                
                for doc in source_docs:
                    source = doc.metadata.get("source", "Unknown")
                    doc_type = doc.metadata.get("document_type", "Unknown")
                    
                    source_info = f"{doc_type} ({source})"
                    if source_info not in unique_sources:
                        unique_sources.add(source_info)
                        sources_text += f"- {source_info}\n"
                
                answer += sources_text
            
            return answer
            
        except Exception as e:
            logger.error(f"Error querying RAG system: {str(e)}")
            return f"Error querying the knowledge base: {str(e)}"
    
    def query_regulations(self, doc_type, topic=None):
        """Query for specific regulations related to document type and topic"""
        if not self.vectorstore or not self.qa_chain or not self.llm:
            return f"RAG system not initialized. Cannot retrieve regulations for {doc_type}."
            
        try:
            # Create specific query for regulations
            if topic:
                query = f"What are the ADGM regulations for {doc_type} regarding {topic}? Please cite specific sections and requirements."
            else:
                query = f"What are the key ADGM regulations for {doc_type}? Please cite specific sections and requirements."
                
            return self.query(query)
        except Exception as e:
            logger.error(f"Error querying regulations: {e}")
            return f"Error retrieving regulations: {str(e)}"