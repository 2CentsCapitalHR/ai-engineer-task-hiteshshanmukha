import os
import json
import requests
import urllib.parse
import time
import re
from bs4 import BeautifulSoup
from pathlib import Path
import tempfile
import logging
from datetime import datetime
import PyPDF2
import docx

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger('adgm_scraper')

class ADGMScraper:
    """Scrapes ADGM websites and documents for regulatory information"""
    
    def __init__(self, output_dir="adgm_extracted_data"):
        self.output_dir = output_dir
        self.documents_dir = os.path.join(output_dir, "documents")
        self.extracted_dir = os.path.join(output_dir, "extracted_text")
        
        # Create directories
        os.makedirs(self.documents_dir, exist_ok=True)
        os.makedirs(self.extracted_dir, exist_ok=True)
        
        # Metadata for all collected documents
        self.document_metadata = []
        
        # Browser-like headers for requests
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8'
        }
    
    def scrape_sources(self, sources):
        """Scrape all provided sources"""
        logger.info(f"Starting to scrape {len(sources)} ADGM data sources")
        
        for idx, (category, doc_type, url) in enumerate(sources, 1):
            logger.info(f"[{idx}/{len(sources)}] Scraping {category} - {doc_type}: {url}")
            
            try:
                if url.lower().endswith(('.pdf', '.docx', '.doc')):
                    self.download_and_process_document(url, category, doc_type)
                else:
                    self.scrape_webpage(url, category, doc_type)
                
                # Be nice to the server
                time.sleep(2)
            except Exception as e:
                logger.error(f"Error scraping {url}: {str(e)}")
        
        # Save metadata
        self.save_metadata()
        
        logger.info(f"Scraping completed. Collected {len(self.document_metadata)} documents")
        return self.document_metadata
    
    def scrape_webpage(self, url, category, doc_type):
        """Scrape a webpage for content and document links"""
        logger.info(f"Scraping webpage: {url}")
        
        try:
            response = requests.get(url, headers=self.headers, timeout=30)
            response.raise_for_status()
            
            # Parse HTML
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Extract and save webpage text
            webpage_text = self.extract_text_from_soup(soup)
            if webpage_text:
                # Create filename from URL
                url_parts = urllib.parse.urlparse(url)
                filename = url_parts.path.strip('/')
                filename = re.sub(r'[^\w\-]', '_', filename)
                if not filename:
                    filename = re.sub(r'[^\w\-]', '_', url_parts.netloc)
                
                text_filename = f"{category}_{doc_type}_{filename}.txt"
                text_path = os.path.join(self.extracted_dir, text_filename)
                
                with open(text_path, 'w', encoding='utf-8') as f:
                    f.write(webpage_text)
                
                # Add to metadata
                self.document_metadata.append({
                    "category": category,
                    "document_type": doc_type,
                    "url": url,
                    "filename": text_filename,
                    "is_webpage": True,
                    "extracted_text_path": text_path,
                    "extraction_date": datetime.now().isoformat()
                })
                
                logger.info(f"Saved webpage text to {text_path}")
            
            # Look for document links to download
            self.extract_and_download_documents(soup, url, category, doc_type)
            
        except Exception as e:
            logger.error(f"Error scraping webpage {url}: {str(e)}")
    
    def extract_and_download_documents(self, soup, base_url, category, doc_type):
        """Extract and download document links from webpage"""
        # Find all links
        links = soup.find_all('a', href=True)
        
        for link in links:
            href = link['href']
            link_text = link.get_text().strip()
            
            # Skip empty or javascript links
            if not href or href.startswith('javascript:') or href == '#':
                continue
            
            # Make absolute URL if needed
            if not href.startswith('http'):
                href = urllib.parse.urljoin(base_url, href)
            
            # Check if it's a document
            if self.is_document_link(href):
                try:
                    logger.info(f"Found document link: {href}")
                    self.download_and_process_document(href, category, doc_type, link_text)
                    time.sleep(1)  # Be nice to the server
                except Exception as e:
                    logger.error(f"Error downloading document {href}: {str(e)}")
    
    def is_document_link(self, url):
        """Check if a URL likely points to a downloadable document"""
        url_lower = url.lower()
        doc_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt', '.txt']
        
        # Check file extension
        if any(url_lower.endswith(ext) for ext in doc_extensions):
            return True
        
        # Check URL keywords suggesting a document
        doc_keywords = ['download', 'document', 'pdf', 'docx', 'template', 'form', 'publication']
        if any(keyword in url_lower for keyword in doc_keywords):
            return True
            
        return False
    
    def download_and_process_document(self, url, category, doc_type, link_text=None):
        """Download document and extract its text"""
        try:
            logger.info(f"Downloading document: {url}")
            
            # Download file
            response = requests.get(url, headers=self.headers, timeout=60, stream=True)
            response.raise_for_status()
            
            # Get filename from URL if not provided in Content-Disposition
            filename = self.get_filename_from_url(url)
            
            # Clean filename to be safe
            filename = re.sub(r'[^\w\-\.]', '_', filename)
            
            # Create path with category prefix
            doc_path = os.path.join(self.documents_dir, f"{category}_{filename}")
            
            # Save document
            with open(doc_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    f.write(chunk)
            
            logger.info(f"Saved document to {doc_path}")
            
            # Extract text based on file type
            extracted_text = None
            if filename.lower().endswith('.pdf'):
                extracted_text = self.extract_text_from_pdf(doc_path)
            elif filename.lower().endswith(('.docx', '.doc')):
                extracted_text = self.extract_text_from_docx(doc_path)
            
            # Save extracted text if successful
            if extracted_text:
                text_filename = os.path.splitext(os.path.basename(doc_path))[0] + ".txt"
                text_path = os.path.join(self.extracted_dir, text_filename)
                
                with open(text_path, 'w', encoding='utf-8') as f:
                    f.write(extracted_text)
                
                logger.info(f"Extracted text to {text_path}")
                
                # Add to metadata
                self.document_metadata.append({
                    "category": category,
                    "document_type": doc_type,
                    "url": url,
                    "original_filename": filename,
                    "saved_path": doc_path,
                    "extracted_text_path": text_path,
                    "link_text": link_text,
                    "extraction_date": datetime.now().isoformat()
                })
        
        except Exception as e:
            logger.error(f"Error processing document {url}: {str(e)}")
    
    def extract_text_from_pdf(self, pdf_path):
        """Extract text content from PDF"""
        text = ""
        try:
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
            return None
    
    def extract_text_from_docx(self, docx_path):
        """Extract text content from DOCX"""
        text = ""
        try:
            doc = docx.Document(docx_path)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
            return None
    
    def extract_text_from_soup(self, soup):
        """Extract meaningful text from HTML"""
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
        
        # Get text
        text = soup.get_text(separator='\n', strip=True)
        
        # Clean up the text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    
    def get_filename_from_url(self, url):
        """Extract filename from URL"""
        parsed_url = urllib.parse.urlparse(url)
        path = parsed_url.path
        
        # Get the last part of the path
        filename = os.path.basename(path)
        
        # If no filename, use a generic one
        if not filename:
            filename = "document"
            
            # Try to add appropriate extension
            if "pdf" in url.lower():
                filename += ".pdf"
            elif "doc" in url.lower():
                filename += ".docx"
            else:
                filename += ".txt"
        
        return filename
    
    def save_metadata(self):
        """Save document metadata to JSON file"""
        metadata_path = os.path.join(self.output_dir, "document_metadata.json")
        
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(self.document_metadata, f, indent=2)
        
        logger.info(f"Saved metadata for {len(self.document_metadata)} documents to {metadata_path}")

def main():
    """Main function to run the scraper"""
    # Define ADGM data sources
    sources = [
        ("Company Formation", "General Incorporation", "https://www.adgm.com/registration-authority/registration-and-incorporation"),
        ("Company Formation", "Resolution for Incorporation", "https://assets.adgm.com/download/assets/adgm-ra-resolution-multiple-incorporate-shareholders-LTD-incorporation-v2.docx/186a12846c3911efa4e6c6223862cd87"),
        ("Company Formation", "Templates", "https://www.adgm.com/setting-up"),
        ("Policy & Guidance", "Templates", "https://www.adgm.com/legal-framework/guidance-and-policy-statements"),
        ("Company Setup", "Checklist", "https://www.adgm.com/documents/registration-authority/registration-and-incorporation/checklist/branch-non-financial-services-20231228.pdf"),
        ("Company Setup", "Private Company Limited", "https://www.adgm.com/documents/registration-authority/registration-and-incorporation/checklist/private-company-limited-by-guarantee-non-financial-services-20231228.pdf"),
        ("Employment", "Standard Contract 2024", "https://assets.adgm.com/download/assets/ADGM+Standard+Employment+Contract+Template+-+ER+2024+(Feb+2025).docx/ee14b252edbe11efa63b12b3a30e5e3a"),
        ("Employment", "Standard Contract 2019", "https://assets.adgm.com/download/assets/ADGM+Standard+Employment+Contract+-+ER+2019+-+Short+Version+(May+2024).docx/33b57a92ecfe11ef97a536cc36767ef8"),
        ("Data Protection", "Policy Template", "https://www.adgm.com/documents/office-of-data-protection/templates/adgm-dpr-2021-appropriate-policy-document.pdf"),
        ("Compliance", "Annual Accounts", "https://www.adgm.com/operating-in-adgm/obligations-of-adgm-registered-entities/annual-filings/annual-accounts"),
        ("Permits", "Application", "https://www.adgm.com/operating-in-adgm/post-registration-services/letters-and-permits"),
        ("Regulatory Guidance", "Incorporation Package", "https://en.adgm.thomsonreuters.com/rulebook/7-company-incorporation-package"),
        ("Regulatory Guidance", "Shareholder Resolution", "https://assets.adgm.com/download/assets/Templates_SHReso_AmendmentArticles-v1-20220107.docx/97120d7c5af911efae4b1e183375c0b2"),
    ]
    
    # Run the scraper
    scraper = ADGMScraper()
    scraper.scrape_sources(sources)

if __name__ == "__main__":
    main()