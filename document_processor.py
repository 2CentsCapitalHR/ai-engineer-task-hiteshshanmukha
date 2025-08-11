# document_processor.py
import os
import docx
import tempfile
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

class DocumentProcessor:
    """Process ADGM legal documents"""
    
    def __init__(self):
        pass
    
    def process_document(self, file_path):
        """Process a document and extract its content"""
        try:
            filename = os.path.basename(file_path)
            
            # Check if file exists and is a DOCX
            if not os.path.exists(file_path):
                return {"error": f"File not found: {file_path}"}
                
            if not file_path.lower().endswith('.docx'):
                return {"error": f"Unsupported file type. Only .docx files are supported."}
            
            # Extract content from DOCX
            doc = docx.Document(file_path)
            content = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                content += para.text + "\n"
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            content += paragraph.text + "\n"
            
            # Extract structured sections and headings
            sections = self._extract_document_sections(doc)
            
            # Return document info
            return {
                "filename": filename,
                "file_path": file_path,
                "content": content,
                "sections": sections,
                "doc_obj": doc  # Store document object for later modification
            }
            
        except Exception as e:
            return {"error": f"Error processing document: {str(e)}"}
    
    def _extract_document_sections(self, doc):
        """Extract document sections based on headings and formatting"""
        sections = []
        current_section = {"title": "General", "content": "", "paragraphs": []}
        
        for i, para in enumerate(doc.paragraphs):
            # Check if this is a heading
            is_heading = False
            if para.style.name.startswith('Heading'):
                is_heading = True
            elif para.runs and para.runs[0].bold and len(para.text.strip()) < 100:
                is_heading = True
            elif para.text.strip().isupper() and len(para.text.strip()) < 100 and para.text.strip():
                is_heading = True
            
            # If this is a heading, start a new section
            if is_heading and para.text.strip():
                # Save previous section if it has content
                if current_section["content"].strip():
                    sections.append(current_section)
                
                # Start new section
                current_section = {
                    "title": para.text.strip(),
                    "content": "",
                    "paragraphs": [i]
                }
            else:
                # Add to current section
                current_section["content"] += para.text + "\n"
                current_section["paragraphs"].append(i)
        
        # Add the last section
        if current_section["content"].strip():
            sections.append(current_section)
        
        return sections
    
    def add_comments(self, doc_info, issues):
        """Add comments to document based on identified issues"""
        try:
            if "error" in doc_info or not issues:
                return None
                
            # Get document object and file info
            doc = doc_info.get("doc_obj")
            filename = doc_info.get("filename")
            
            if not doc:
                return None
            
            # Create a map of sections to paragraph indices
            section_map = {}
            for section in doc_info.get("sections", []):
                section_title = section.get("title")
                for para_index in section.get("paragraphs", []):
                    section_map[para_index] = section_title
            
            # Create a new temp file for the marked-up document
            temp_dir = tempfile.mkdtemp()
            marked_file_path = os.path.join(temp_dir, f"REVIEWED_{filename}")
            
            # Create a tracking paragraph to list all issues
            summary_para = doc.add_paragraph()
            summary_para.add_run("DOCUMENT REVIEW SUMMARY\n").bold = True
            summary_para.add_run(f"Total issues found: {len(issues)}\n\n")
            
            # Group issues by section for better organization
            issues_by_section = {}
            for issue in issues:
                section = issue.get("section", "General")
                if section not in issues_by_section:
                    issues_by_section[section] = []
                issues_by_section[section].append(issue)
            
            # Add a summary of issues by section
            for section, section_issues in issues_by_section.items():
                section_run = summary_para.add_run(f"{section}: {len(section_issues)} issues\n")
                section_run.bold = True
                
                for issue in section_issues:
                    severity = issue.get("severity", "Medium")
                    issue_text = issue.get("issue", "")
                    
                    # Add color-coded severity
                    if severity == "High":
                        severity_run = summary_para.add_run("HIGH: ")
                        severity_run.font.color.rgb = RGBColor(255, 0, 0)
                    elif severity == "Medium":
                        severity_run = summary_para.add_run("MEDIUM: ")
                        severity_run.font.color.rgb = RGBColor(255, 165, 0)
                    else:
                        severity_run = summary_para.add_run("LOW: ")
                        severity_run.font.color.rgb = RGBColor(0, 0, 255)
                    
                    # Add issue text
                    summary_para.add_run(f"{issue_text}\n")
            
            # Add a page break after summary
            doc.add_page_break()
            
            # Track which paragraphs we've already commented on
            commented_paragraphs = set()
            
            # Apply comments based on issues
            for issue in issues:
                section_name = issue.get("section", "General")
                issue_text = issue.get("issue", "")
                severity = issue.get("severity", "Medium")
                suggestion = issue.get("suggestion", "")
                regulation = issue.get("regulation", "ADGM Regulations")
                
                # Find paragraphs related to this section
                para_indices = self._find_section_paragraphs(doc_info, section_name)
                
                if para_indices:
                    # Choose the first paragraph we haven't commented on yet, if possible
                    target_index = None
                    for idx in para_indices:
                        if idx not in commented_paragraphs:
                            target_index = idx
                            commented_paragraphs.add(idx)
                            break
                    
                    # If all paragraphs already have comments, just use the first one
                    if target_index is None and para_indices:
                        target_index = para_indices[0]
                    
                    if target_index is not None:
                        para = doc.paragraphs[target_index]
                        
                        # Highlight text based on severity
                        if severity == "High":
                            self._highlight_paragraph(para, WD_COLOR_INDEX.RED)
                        elif severity == "Medium":
                            self._highlight_paragraph(para, WD_COLOR_INDEX.YELLOW)
                        else:
                            self._highlight_paragraph(para, WD_COLOR_INDEX.TURQUOISE)
                        
                        # Create a detailed comment paragraph
                        comment_para = doc.add_paragraph()
                        
                        # Add an icon to visually separate the comment
                        if severity == "High":
                            icon = "⚠️ "  # Warning icon for high severity
                        elif severity == "Medium":
                            icon = "ℹ️ "  # Info icon for medium severity
                        else:
                            icon = "📝 "  # Note icon for low severity
                        
                        # Create the comment header
                        comment_header = comment_para.add_run(f"{icon}ISSUE ({severity} severity): ")
                        comment_header.bold = True
                        comment_header.font.color.rgb = self._get_color_for_severity(severity)
                        
                        # Add the issue text
                        comment_para.add_run(f"{issue_text}\n")
                        
                        # Add suggestion if available
                        if suggestion:
                            suggestion_run = comment_para.add_run("SUGGESTION: ")
                            suggestion_run.bold = True
                            comment_para.add_run(f"{suggestion}\n")
                        
                        # Add regulation reference
                        if regulation:
                            regulation_run = comment_para.add_run("REGULATION: ")
                            regulation_run.bold = True
                            regulation_run.italic = True
                            comment_para.add_run(f"{regulation}\n")
                        
                        # Add a separator line
                        comment_para.add_run("________________________________________\n")
            
            # Save the marked-up document
            doc.save(marked_file_path)
            return marked_file_path
            
        except Exception as e:
            print(f"Error adding comments to document: {str(e)}")
            return None
    
    def _find_section_paragraphs(self, doc_info, section_name):
        """Find paragraph indices that belong to a specific section"""
        # First try to find an exact section match
        for section in doc_info.get("sections", []):
            if section.get("title") == section_name:
                return section.get("paragraphs", [])
        
        # If not found, try a fuzzy match
        for section in doc_info.get("sections", []):
            if section_name.lower() in section.get("title", "").lower():
                return section.get("paragraphs", [])
            
            # Check if section content contains keywords related to the section name
            content_lower = section.get("content", "").lower()
            section_keywords = section_name.lower().split()
            
            matches = 0
            for keyword in section_keywords:
                if len(keyword) > 3 and keyword in content_lower:  # Only check substantial keywords
                    matches += 1
            
            # If most keywords match, consider it a match
            if matches >= max(1, len(section_keywords) // 2):
                return section.get("paragraphs", [])
        
        # If still not found, return first few paragraphs as a fallback
        doc = doc_info.get("doc_obj")
        if doc and len(doc.paragraphs) > 0:
            return [0]  # Return first paragraph
            
        return []
    
    def _highlight_paragraph(self, paragraph, color):
        """Highlight text in a paragraph"""
        for run in paragraph.runs:
            run.font.highlight_color = color
    
    def _get_color_for_severity(self, severity):
        """Get RGB color based on severity"""
        if severity == "High":
            return RGBColor(255, 0, 0)  # Red
        elif severity == "Medium":
            return RGBColor(255, 165, 0)  # Orange
        else:
            return RGBColor(0, 0, 255)  # Blue